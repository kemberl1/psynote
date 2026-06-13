"""Извлечение СЫРОГО текста из файлов корпуса (docs/03 §5, уровень 1 docs/04).

Поддерживаемые форматы:
  - .docx — python-docx (zip + XML);
  - .odt  — odfpy (zip + content.xml);
  - .doc  — старый бинарный формат: пытаемся через `antiword`/LibreOffice headless,
            если инструмента нет в образе — корректно пропускаем с логом;
  - .xlsx — openpyxl (листы назначений). По docs/03 §4 фокус Этапа 3 — ДНЕВНИКИ
            (.docx/.odt). Таблицы .xlsx по умолчанию НЕ индексируются вместе с
            дневниками (другой регистр текста), но экстрактор реализован и может
            быть включён флагом — см. ingest.py.

ВАЖНО (приватность): здесь извлекается СЫРОЙ текст с возможными ПДн. Этот текст
НИКОГДА не логируется и НИКОГДА не пишется в Qdrant напрямую — он идёт ТОЛЬКО в
анонимайзер gateway (см. anonymizer_client.py, ingest.py). Имена файлов с ФИО
тоже НЕ попадают в payload (см. ingest.py — используется хэш-идентификатор).
"""

from __future__ import annotations

import logging
import shutil
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)

# Форматы, которые умеем извлекать текстом.
SUPPORTED_TEXT_SUFFIXES = {".docx", ".odt", ".doc"}
SUPPORTED_TABLE_SUFFIXES = {".xlsx"}
SUPPORTED_SUFFIXES = SUPPORTED_TEXT_SUFFIXES | SUPPORTED_TABLE_SUFFIXES


class ExtractionError(Exception):
    """Не удалось извлечь текст (битый файл / нет инструмента / формат)."""


def extract_text(path: Path) -> str:
    """Извлечь сырой текст из файла по расширению.

    Бросает ExtractionError при невозможности (вызывающий код логирует путь
    БЕЗ содержимого и пропускает файл — устойчивость к битым файлам).
    """
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".odt":
        return _extract_odt(path)
    if suffix == ".doc":
        return _extract_doc(path)
    if suffix == ".xlsx":
        return _extract_xlsx(path)
    raise ExtractionError(f"неподдерживаемый формат: {suffix}")


# ─── .docx ──────────────────────────────────────────────────────────────────
def _extract_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - зависит от образа
        raise ExtractionError("python-docx не установлен") from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # битый zip/xml
        raise ExtractionError(
            f"не удалось открыть .docx: {type(exc).__name__}") from exc

    parts: list[str] = []
    # Абзацы основного тела.
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Таблицы (в дневниках встречаются как разметка) — построчно.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ─── .odt ───────────────────────────────────────────────────────────────────
def _extract_odt(path: Path) -> str:
    try:
        from odf import teletype, text  # type: ignore
        from odf.opendocument import load  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("odfpy не установлен") from exc

    try:
        doc = load(str(path))
    except Exception as exc:
        raise ExtractionError(
            f"не удалось открыть .odt: {type(exc).__name__}") from exc

    parts: list[str] = []
    for elem in doc.getElementsByType(text.P):
        line = teletype.extractText(elem).strip()
        if line:
            parts.append(line)
    return "\n".join(parts)


# ─── .doc (старый бинарный формат) ────────────────────────────────────────────
def _extract_doc(path: Path) -> str:
    """Пытаемся antiword → LibreOffice headless. Иначе ExtractionError (пропуск)."""
    # 1) antiword — лёгкий и точный для .doc.
    if shutil.which("antiword"):
        try:
            out = subprocess.run(
                ["antiword", str(path)],
                capture_output=True, timeout=60, check=True,
            )
            return out.stdout.decode("utf-8", errors="replace").strip()
        except Exception as exc:
            logger.warning("antiword не справился с .doc (%s), пробую LibreOffice",
                           type(exc).__name__)

    # 2) LibreOffice headless: конвертация в .txt во временную папку.
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            with tempfile.TemporaryDirectory() as tmp:
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "txt:Text",
                     "--outdir", tmp, str(path)],
                    capture_output=True, timeout=120, check=True,
                )
                produced = next(Path(tmp).glob("*.txt"), None)
                if produced and produced.exists():
                    return produced.read_text(encoding="utf-8", errors="replace").strip()
        except Exception as exc:
            logger.warning(
                "LibreOffice не справился с .doc (%s)", type(exc).__name__)

    raise ExtractionError(
        ".doc не извлечён: нет antiword/LibreOffice в образе (файл пропущен)")


# ─── .xlsx ────────────────────────────────────────────────────────────────────
def _extract_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("openpyxl не установлен") from exc

    try:
        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    except Exception as exc:
        raise ExtractionError(
            f"не удалось открыть .xlsx: {type(exc).__name__}") from exc

    parts: list[str] = []
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip()
                         for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
    finally:
        wb.close()
    return "\n".join(parts)
